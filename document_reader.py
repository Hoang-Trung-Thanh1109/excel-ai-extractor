"""Read PDF and DOCX documents into plain text context."""

from __future__ import annotations

from pathlib import Path
from typing import Iterable


class DocumentReader:
    def read_documents(self, file_paths: Iterable[str]) -> str:
        blocks: list[str] = []
        for raw_path in file_paths:
            path = Path(raw_path)
            if not path.exists():
                continue

            if path.suffix.lower() == ".pdf":
                text = self._read_pdf(path)
            elif path.suffix.lower() == ".docx":
                text = self._read_docx(path)
            elif path.suffix.lower() in {".txt", ".md"}:
                text = path.read_text(encoding="utf-8", errors="ignore")
            else:
                text = path.read_text(encoding="utf-8", errors="ignore")

            if text.strip():
                blocks.append(f"FILE: {path.name}\n{text.strip()}")

        return "\n\n".join(blocks).strip()

    def _read_pdf(self, path: Path) -> str:
        try:
            import pdfplumber
        except ImportError as exc:
            raise RuntimeError(
                "Thieu goi 'pdfplumber'. Hay cai dat trong requirements.txt truoc khi chay."
            ) from exc

        chunks: list[str] = []
        with pdfplumber.open(str(path)) as pdf:
            for page_index, page in enumerate(pdf.pages, start=1):
                page_text = page.extract_text() or ""
                if page_text.strip():
                    chunks.append(f"[Page {page_index}]\n{page_text.strip()}")
        return "\n\n".join(chunks)

    def _read_docx(self, path: Path) -> str:
        try:
            from docx import Document
        except ImportError as exc:
            raise RuntimeError(
                "Thieu goi 'python-docx'. Hay cai dat trong requirements.txt truoc khi chay."
            ) from exc

        doc = Document(str(path))
        paragraphs = [paragraph.text.strip() for paragraph in doc.paragraphs if paragraph.text.strip()]

        tables: list[str] = []
        for table_index, table in enumerate(doc.tables, start=1):
            rows = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                rows.append(" | ".join(cells))
            if rows:
                tables.append(f"[Table {table_index}]\n" + "\n".join(rows))

        return "\n\n".join(paragraphs + tables)
